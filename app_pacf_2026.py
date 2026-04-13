import html
import io
import re
import unicodedata
from typing import Dict, Optional

import numpy as np
import pandas as pd
import streamlit as st
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ============================================================
# CONFIGURACIÓN VISUAL
# ============================================================
st.set_page_config(page_title="PACF 2026 - Anexo I", layout="wide")

st.markdown("""
<style>
    .block-container {padding-top: 1.2rem; padding-bottom: 2rem; max-width: 1400px;}
    h1, h2, h3 {font-family: "Segoe UI", sans-serif;}
    .main-title {
        text-align:center;
        font-weight:700;
        font-size: 1.75rem;
        line-height: 1.25;
        margin-bottom: 0.35rem;
        white-space: normal;
        word-break: normal;
        overflow-wrap: anywhere;
    }
    .sub-title {
        text-align:center;
        font-size: 1rem;
        line-height: 1.35;
        color:#4b5563;
        margin-bottom: 1.2rem;
        white-space: normal;
        overflow-wrap: anywhere;
    }
    .section-card {
        border:1px solid #d1d5db;
        border-radius:10px;
        padding:0.8rem 1rem;
        background:#f9fafb;
        margin-bottom:0.8rem;
    }
    .kpi {
        border:1px solid #d1d5db;
        border-radius:10px;
        padding:0.8rem;
        background:white;
        text-align:center;
    }
    .kpi-title {
        font-size:0.9rem;
        color:#6b7280;
        margin-bottom:0.2rem;
    }
    .kpi-value {
        font-size:1.4rem;
        font-weight:700;
        color:#111827;
    }
    .method-note {
        border-left:4px solid #9ca3af;
        background:#f3f4f6;
        padding:0.8rem 1rem;
        border-radius:6px;
        font-size:0.95rem;
    }
    .badge {
        display:inline-block;
        padding:0.18rem 0.55rem;
        border-radius:999px;
        font-size:0.82rem;
        font-weight:600;
        border:1px solid #d1d5db;
    }
    @media (max-width: 900px) {
        .main-title {font-size: 1.35rem;}
        .sub-title {font-size: 0.95rem;}
    }
</style>
""", unsafe_allow_html=True)

st.markdown('<div class="main-title">ANEXO I · EVALUACIÓN ESTADÍSTICA DE LOS RIESGOS</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-title">Versión visual PACF 2026 · Probabilidad, impacto financiero y matriz de riesgo por secciones</div>', unsafe_allow_html=True)


# ============================================================
# UTILIDADES DE FORMATO
# ============================================================
def fmt_es_num(x, dec=2):
    if pd.isna(x):
        return ""
    s = f"{float(x):,.{dec}f}"
    return s.replace(",", "X").replace(".", ",").replace("X", ".")


def fmt_es_pct(x, dec=2):
    if pd.isna(x):
        return ""
    return f"{fmt_es_num(x, dec)} %"


def fmt_es_eur(x, dec=2):
    if pd.isna(x):
        return ""
    return f"{fmt_es_num(x, dec)} €"


def sanitize_columns(df):
    df = df.copy()
    df.columns = [str(c).strip() for c in df.columns]
    return df


def leer_excel(uploaded_file):
    name = getattr(uploaded_file, "name", "").lower()
    engines = ["xlrd"] if name.endswith(".xls") else ["openpyxl"]
    engines.append(None)

    last_error = None
    for engine in engines:
        uploaded_file.seek(0)
        try:
            if engine is None:
                return pd.read_excel(uploaded_file)
            return pd.read_excel(uploaded_file, engine=engine)
        except Exception as exc:
            last_error = exc

    raise ValueError(f"No se ha podido leer el Excel cargado: {last_error}")


def parse_numeric_text(value):
    if pd.isna(value):
        return 0

    text = str(value).strip()
    if not text:
        return 0

    negative = text.startswith("(") and text.endswith(")")
    text = re.sub(r"[^\d,\.\-]", "", text)
    negative = negative or text.startswith("-")
    text = text.replace("-", "")

    if not text:
        return 0

    if "," in text and "." in text:
        decimal_sep = "," if text.rfind(",") > text.rfind(".") else "."
    elif "," in text:
        decimal_sep = ","
    elif "." in text:
        decimals = text.rsplit(".", 1)[1]
        decimal_sep = "." if 0 < len(decimals) <= 2 else None
    else:
        decimal_sep = None

    if decimal_sep:
        thousands_sep = "." if decimal_sep == "," else ","
        text = text.replace(thousands_sep, "")
        text = text.replace(decimal_sep, ".")
    else:
        text = text.replace(".", "").replace(",", "")

    try:
        number = float(text)
    except ValueError:
        return 0
    return -number if negative and number > 0 else number


def to_numeric_safe(series):
    if pd.api.types.is_numeric_dtype(series):
        return pd.to_numeric(series, errors="coerce").fillna(0)
    return series.map(parse_numeric_text).fillna(0)


# ============================================================
# METODOLOGÍA
# ============================================================
PROB_BANDS = [
    ("Raro", 0, 10),
    ("Improbable", 10, 20),
    ("Posible", 20, 40),
    ("Probable", 40, 80),
    ("Esperado", 80, float("inf")),
]

IMPACT_BANDS = [
    ("Muy bajo", 0, 0.1),
    ("Bajo", 0.1, 2),
    ("Medio", 2, 10),
    ("Alto", 10, 25),
    ("Muy alto", 25, float("inf")),
]

RISK_MATRIX = {
    ("Raro", "Muy bajo"): "Bajo",
    ("Raro", "Bajo"): "Bajo",
    ("Raro", "Medio"): "Bajo",
    ("Raro", "Alto"): "Bajo",
    ("Raro", "Muy alto"): "Medio",
    ("Improbable", "Muy bajo"): "Bajo",
    ("Improbable", "Bajo"): "Bajo",
    ("Improbable", "Medio"): "Bajo",
    ("Improbable", "Alto"): "Medio",
    ("Improbable", "Muy alto"): "Medio",
    ("Posible", "Muy bajo"): "Bajo",
    ("Posible", "Bajo"): "Bajo",
    ("Posible", "Medio"): "Medio",
    ("Posible", "Alto"): "Medio",
    ("Posible", "Muy alto"): "Alto",
    ("Probable", "Muy bajo"): "Bajo",
    ("Probable", "Bajo"): "Medio",
    ("Probable", "Medio"): "Medio",
    ("Probable", "Alto"): "Alto",
    ("Probable", "Muy alto"): "Alto",
    ("Esperado", "Muy bajo"): "Medio",
    ("Esperado", "Bajo"): "Medio",
    ("Esperado", "Medio"): "Alto",
    ("Esperado", "Alto"): "Alto",
    ("Esperado", "Muy alto"): "Alto",
}

PROB_ORDER = ["Raro", "Improbable", "Posible", "Probable", "Esperado"]
IMPACT_ORDER = ["Muy bajo", "Bajo", "Medio", "Alto", "Muy alto"]

EXCLUDED_STATES_DEFAULT = {"Borrador", "Pendiente visto bueno del área", "Anulado"}


def classify_band(value, bands):
    try:
        v = float(value)
    except Exception:
        return ""
    for label, low, high in bands:
        if low <= v < high:
            return label
    return bands[-1][0]


def badge_html(text):
    palette = {
        "Raro": "#e5e7eb", "Improbable": "#dbeafe", "Posible": "#fef3c7",
        "Probable": "#fed7aa", "Esperado": "#fecaca",
        "Muy bajo": "#f3f4f6", "Bajo": "#dbeafe", "Medio": "#fde68a",
        "Alto": "#fdba74", "Muy alto": "#fca5a5",
        "Bajo_riesgo": "#dcfce7", "Medio_riesgo": "#fde68a", "Alto_riesgo": "#fecaca",
        "No disponible": "#f3f4f6"
    }
    key = text
    if text == "Bajo":
        key = "Bajo_riesgo"
    elif text == "Medio":
        key = "Medio_riesgo"
    elif text == "Alto":
        key = "Alto_riesgo"
    color = palette.get(key, "#f3f4f6")
    escaped_text = html.escape("" if text is None else str(text))
    return f'<span class="badge" style="background:{color};">{escaped_text}</span>'


# ============================================================
# MAPEO DE SECCIONES
# ============================================================
def normalizar_nombre_columna(col: str) -> str:
    col = str(col).strip().lower()
    col = "".join(
        char for char in unicodedata.normalize("NFKD", col)
        if not unicodedata.combining(char)
    )
    col = col.replace("º", "").replace("ª", "")
    col = col.replace(".", " ").replace("-", " ").replace("_", " ").replace("/", " ")
    col = " ".join(col.split())
    return col


def guess_column(columns, aliases, contains_all=None):
    normalized = {col: normalizar_nombre_columna(col) for col in columns if col}
    alias_keys = {normalizar_nombre_columna(alias) for alias in aliases}

    exact = next((col for col, norm in normalized.items() if norm in alias_keys), None)
    if exact:
        return exact

    if contains_all:
        required = [normalizar_nombre_columna(token) for token in contains_all]
        return next(
            (col for col, norm in normalized.items() if all(token in norm for token in required)),
            None,
        )
    return None


def selectbox_column(label, cols, aliases, key, contains_all=None):
    guessed = guess_column(cols, aliases, contains_all=contains_all)
    index = cols.index(guessed) if guessed in cols else 0
    return st.selectbox(label, cols, index=index, key=key)

def preparar_mapeo_secciones(df_map: pd.DataFrame) -> pd.DataFrame:
    dfm = sanitize_columns(df_map)

    aliases_seccion = {
        "seccion", "codigo seccion", "cod seccion", "seccion informe",
        "seccion del informe", "codigo", "cod", "id seccion"
    }
    aliases_descripcion = {
        "descripcion", "denominacion", "nombre", "detalle", "concepto",
        "descripcion seccion", "nombre seccion", "seccion descripcion"
    }

    normalized = {col: normalizar_nombre_columna(col) for col in dfm.columns}

    col_seccion = next((orig for orig, norm in normalized.items() if norm in aliases_seccion), None)
    col_descripcion = next((orig for orig, norm in normalized.items() if norm in aliases_descripcion), None)

    # Búsqueda más tolerante por palabras clave
    if col_seccion is None:
        col_seccion = next(
            (orig for orig, norm in normalized.items() if "seccion" in norm and ("codigo" in norm or "informe" in norm or norm == "seccion")),
            None
        )
    if col_descripcion is None:
        col_descripcion = next(
            (orig for orig, norm in normalized.items() if "descripcion" in norm or "denominacion" in norm or "nombre" in norm),
            None
        )

    if col_seccion is None or col_descripcion is None:
        raise ValueError(
            "No se han podido identificar las columnas del mapeo. "
            "La hoja debe incluir una columna de sección y otra de descripción. "
            f"Columnas detectadas: {dfm.columns.tolist()}"
        )

    dfm = dfm[[col_seccion, col_descripcion]].copy()
    dfm.columns = ["Sección", "Descripción"]
    dfm["Sección"] = dfm["Sección"].astype(str).str.strip().str[:4]
    dfm["Descripción"] = dfm["Descripción"].astype(str).str.strip()
    dfm = dfm[~dfm["Sección"].str.lower().isin(["", "nan", "none"])].copy()
    dfm = dfm.drop_duplicates(subset=["Sección"], keep="first").reset_index(drop=True)
    dfm["Orden_mapeo"] = range(1, len(dfm) + 1)
    return dfm


def aplicar_mapeo(df: pd.DataFrame, df_map: Optional[pd.DataFrame]) -> pd.DataFrame:
    out = df.copy()
    if df_map is None or df_map.empty:
        if "Descripción" not in out.columns:
            out["Descripción"] = ""
        out["Orden_mapeo"] = 999999
        return out

    out = out.drop(columns=[c for c in ["Descripción", "Orden_mapeo"] if c in out.columns])
    out = out.merge(df_map, on="Sección", how="left")
    out["Descripción"] = out["Descripción"].fillna("")
    out["Orden_mapeo"] = out["Orden_mapeo"].fillna(999999)
    return out


def ordenar_por_mapeo(df: pd.DataFrame) -> pd.DataFrame:
    cols = df.columns.tolist()
    sort_cols = [c for c in ["Orden_mapeo", "Sección"] if c in df.columns]
    return df.sort_values(sort_cols).reset_index(drop=True)[cols]


# ============================================================
# NORMALIZACIÓN
# ============================================================
def get_section_from_expediente(series):
    return series.astype(str).str.strip().str[:4]


def normalize_year_df(df_raw: pd.DataFrame, year: int, column_map: Dict[str, Optional[str]]) -> pd.DataFrame:
    df = sanitize_columns(df_raw)

    missing = []
    for logical_name in ["expediente", "importe", "desfavorables"]:
        col = column_map.get(logical_name)
        if not col or col not in df.columns:
            missing.append(logical_name)
    if missing:
        raise ValueError(f"Faltan columnas obligatorias mapeadas: {missing}")

    out = pd.DataFrame()
    out["Año"] = year
    out["Nº Expediente FLP"] = df[column_map["expediente"]].astype(str).str.strip()
    out["Sección"] = get_section_from_expediente(out["Nº Expediente FLP"])
    out["Importe"] = to_numeric_safe(df[column_map["importe"]])
    out["Número de Informes Desfavorables"] = to_numeric_safe(df[column_map["desfavorables"]])

    fav_col = column_map.get("favorables")
    out["Número de Informes Favorables"] = to_numeric_safe(df[fav_col]) if fav_col and fav_col in df.columns else 0

    estado_col = column_map.get("estado")
    out["Estado"] = df[estado_col].astype(str).str.strip() if estado_col and estado_col in df.columns else ""

    fase_col = column_map.get("fase")
    out["Fase del Gasto"] = df[fase_col].astype(str).str.strip() if fase_col and fase_col in df.columns else ""

    return out


def depurar_df(df: pd.DataFrame, exclude_states=None):
    exclude_states = exclude_states or EXCLUDED_STATES_DEFAULT
    if "Estado" not in df.columns:
        return df.copy(), pd.DataFrame(columns=list(df.columns) + ["Motivo exclusión"])

    excluded_keys = {normalizar_nombre_columna(value) for value in exclude_states}
    state_keys = df["Estado"].fillna("").map(normalizar_nombre_columna)
    mask = state_keys.isin(excluded_keys)
    excluidos = df[mask].copy()
    if not excluidos.empty:
        excluidos["Motivo exclusión"] = "Estado excluido"
    return df[~mask].copy(), excluidos


# ============================================================
# CÁLCULO
# ============================================================
def calcular_probabilidad_anual(df_year: pd.DataFrame, df_map: Optional[pd.DataFrame] = None) -> pd.DataFrame:
    df = df_year.copy()
    df["Its_fila"] = df["Número de Informes Favorables"] + df["Número de Informes Desfavorables"]
    it_total = df["Its_fila"].sum()

    grouped = (
        df.groupby("Sección", dropna=False)
        .agg(
            **{
                "Id*s": ("Número de Informes Desfavorables", "sum"),
                "Its": ("Its_fila", "sum"),
                "Expedientes": ("Nº Expediente FLP", "nunique"),
                "Ms": ("Importe", "sum"),
            }
        )
        .reset_index()
    )
    grouped["It"] = it_total
    grouped["P1 (Id*s / Its)"] = np.where(grouped["Its"] > 0, grouped["Id*s"] / grouped["Its"] * 100, 0)
    grouped["P2 (Its / It)"] = np.where(it_total > 0, grouped["Its"] / it_total * 100, 0)
    grouped["Ps (%)"] = ((grouped["P1 (Id*s / Its)"] * 65) + (grouped["P2 (Its / It)"] * 35)) / 100
    grouped["Nivel de probabilidad"] = grouped["Ps (%)"].apply(lambda x: classify_band(x, PROB_BANDS))

    grouped = aplicar_mapeo(grouped, df_map)
    return ordenar_por_mapeo(grouped)


def calcular_impacto_anual(df_year: pd.DataFrame, df_map: Optional[pd.DataFrame] = None, modo_ms: str = "Solo peticiones con 1 o más desfavorables") -> pd.DataFrame:
    df = df_year.copy()
    m_total = df["Importe"].sum()

    if modo_ms == "Todos los informes válidos":
        df_base = df.copy()
        descripcion_modo = "Financiero (todos los informes válidos)"
    else:
        df_base = df[df["Número de Informes Desfavorables"] >= 1].copy()
        descripcion_modo = "Financiero (solo peticiones con ≥1 desfavorable)"

    grouped = (
        df_base.groupby("Sección", dropna=False)
        .agg(
            Ms=("Importe", "sum"),
            Expedientes=("Nº Expediente FLP", "nunique")
        )
        .reset_index()
    )

    # Incluir todas las secciones del ejercicio, aunque queden con Ms=0
    todas_secciones = pd.DataFrame({"Sección": sorted(df["Sección"].dropna().astype(str).unique().tolist())})
    grouped = todas_secciones.merge(grouped, on="Sección", how="left")
    grouped["Ms"] = grouped["Ms"].fillna(0)
    grouped["Expedientes"] = grouped["Expedientes"].fillna(0)

    grouped["M"] = m_total
    grouped["Is (%)"] = np.where(m_total > 0, grouped["Ms"] / m_total * 100, 0)
    grouped["Nivel de impacto"] = grouped["Is (%)"].apply(lambda x: classify_band(x, IMPACT_BANDS))
    grouped["Severidad"] = "No disponible"
    grouped["Modo impacto"] = descripcion_modo

    grouped = aplicar_mapeo(grouped, df_map)
    return ordenar_por_mapeo(grouped)


def consolidar_probabilidad(prob_tables: Dict[int, pd.DataFrame], df_map: Optional[pd.DataFrame], pesos_probabilidad=None) -> pd.DataFrame:
    years = sorted(prob_tables.keys())
    if not years:
        return pd.DataFrame()

    base_sections = pd.concat([t[["Sección"]] for t in prob_tables.values()]).drop_duplicates()
    base_sections = aplicar_mapeo(base_sections, df_map)
    out = ordenar_por_mapeo(base_sections)

    if pesos_probabilidad is None:
        pesos_probabilidad = [20.0, 30.0, 50.0]

    if len(years) == 1:
        raw = [100.0]
        weights = {years[0]: 1.0}
        etiqueta_pesos = "100"
    elif len(years) == 2:
        raw = [pesos_probabilidad[1], pesos_probabilidad[2]]
        total = sum(raw)
        if total == 0:
            raw = [40.0, 60.0]
            total = 100.0
        norm = [v / total for v in raw]
        weights = {years[0]: norm[0], years[1]: norm[1]}
        etiqueta_pesos = f"{fmt_es_num(raw[0],0)}/{fmt_es_num(raw[1],0)}"
    else:
        raw = pesos_probabilidad[:3]
        total = sum(raw)
        if total == 0:
            raw = [20.0, 30.0, 50.0]
            total = 100.0
        norm = [v / total for v in raw]
        weights = {years[0]: norm[0], years[1]: norm[1], years[2]: norm[2]}
        etiqueta_pesos = f"{fmt_es_num(raw[0],0)}/{fmt_es_num(raw[1],0)}/{fmt_es_num(raw[2],0)}"

    weighted_sum = 0
    for y in years:
        tmp = prob_tables[y][["Sección", "Ps (%)"]].rename(columns={"Ps (%)": f"Ps {y}"})
        out = out.merge(tmp, on="Sección", how="left")
        weighted_sum = weighted_sum + out[f"Ps {y}"].fillna(0) * weights[y]

    nombre_col_media = f"Media ponderada ({etiqueta_pesos})"
    out[nombre_col_media] = weighted_sum
    out["Media ponderada probabilidad"] = weighted_sum
    out["Nivel de probabilidad"] = out["Media ponderada probabilidad"].apply(lambda x: classify_band(x, PROB_BANDS))
    out["Pesos aplicados"] = etiqueta_pesos
    return ordenar_por_mapeo(out)


def consolidar_impacto(impact_tables: Dict[int, pd.DataFrame], df_map: Optional[pd.DataFrame]) -> pd.DataFrame:
    years = sorted(impact_tables.keys())
    if not years:
        return pd.DataFrame()

    base_sections = pd.concat([t[["Sección"]] for t in impact_tables.values()]).drop_duplicates()
    base_sections = aplicar_mapeo(base_sections, df_map)
    out = ordenar_por_mapeo(base_sections)

    year_cols = []
    for y in years:
        tmp = impact_tables[y][["Sección", "Is (%)"]].rename(columns={"Is (%)": f"Is {y}"})
        out = out.merge(tmp, on="Sección", how="left")
        year_cols.append(f"Is {y}")

    out["Media trienal impacto"] = out[year_cols].fillna(0).mean(axis=1)
    out["Nivel de impacto"] = out["Media trienal impacto"].apply(lambda x: classify_band(x, IMPACT_BANDS))
    out["Severidad"] = "No disponible 2023-2025"
    out["Modo impacto"] = "Financiero"
    return ordenar_por_mapeo(out)


def construir_matriz_final(prob_final: pd.DataFrame, impact_final: pd.DataFrame) -> pd.DataFrame:
    # Usar siempre la columna fija si existe; si no, localizar la dinámica y normalizarla
    if "Media ponderada probabilidad" in prob_final.columns:
        col_media_prob = "Media ponderada probabilidad"
    else:
        col_media_prob = next((c for c in prob_final.columns if c.startswith("Media ponderada")), None)

    if col_media_prob is None:
        raise ValueError(
            f"No se ha encontrado ninguna columna de media ponderada en prob_final. "
            f"Columnas disponibles: {prob_final.columns.tolist()}"
        )

    base_prob = prob_final[["Sección", "Descripción", "Orden_mapeo", col_media_prob, "Nivel de probabilidad"]].copy()
    if col_media_prob != "Media ponderada probabilidad":
        base_prob = base_prob.rename(columns={col_media_prob: "Media ponderada probabilidad"})

    out = base_prob.merge(
        impact_final[["Sección", "Media trienal impacto", "Nivel de impacto", "Modo impacto"]],
        on="Sección",
        how="outer"
    )
    out["Descripción"] = out["Descripción"].fillna("")
    out["Orden_mapeo"] = out["Orden_mapeo"].fillna(999999)
    out["Nivel de riesgo"] = out.apply(
        lambda r: RISK_MATRIX.get((r.get("Nivel de probabilidad", ""), r.get("Nivel de impacto", "")), ""),
        axis=1
    )
    return ordenar_por_mapeo(out)


def make_unique_sheet_name(name: str, used_names: set) -> str:
    base = re.sub(r"[\[\]\:\*\?\/\\]", "_", str(name).strip()).strip("'")[:31] or "Hoja"
    sheet_name = base
    counter = 2
    while sheet_name.lower() in used_names:
        suffix = f"_{counter}"
        sheet_name = f"{base[:31 - len(suffix)]}{suffix}"
        counter += 1
    used_names.add(sheet_name.lower())
    return sheet_name


def build_excel_export(frames: Dict[str, pd.DataFrame]) -> bytes:
    output = io.BytesIO()
    used_sheet_names = set()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        for name, df in frames.items():
            sheet_name = make_unique_sheet_name(name, used_sheet_names)
            df.to_excel(writer, index=False, sheet_name=sheet_name)
    output.seek(0)
    return output.getvalue()


# ============================================================
# TABLAS VISUALES
# ============================================================
def display_visual_table(df: pd.DataFrame, title: str, percent_cols=None, euro_cols=None, int_cols=None, badge_cols=None):
    st.markdown(f"### {title}")
    if df.empty:
        st.info("No hay datos.")
        return

    percent_cols = percent_cols or []
    euro_cols = euro_cols or []
    int_cols = int_cols or []
    badge_cols = badge_cols or []

    show = df.copy()
    for c in percent_cols:
        if c in show.columns:
            show[c] = show[c].apply(fmt_es_pct)
    for c in euro_cols:
        if c in show.columns:
            show[c] = show[c].apply(fmt_es_eur)
    for c in int_cols:
        if c in show.columns:
            show[c] = show[c].apply(lambda v: fmt_es_num(v, 0))
    for c in show.columns:
        if c not in badge_cols:
            show[c] = show[c].apply(lambda v: "" if pd.isna(v) else html.escape(str(v)))
    for c in badge_cols:
        if c in show.columns:
            show[c] = show[c].apply(badge_html)

    html_table = show.to_html(index=False, escape=False)
    st.markdown(html_table, unsafe_allow_html=True)


def display_matrix_grid(df_matriz: pd.DataFrame):
    st.markdown("### Matriz de nivel de riesgo")
    header_cols = st.columns([1.3, 1, 1, 1, 1, 1])
    header_cols[0].markdown("**PROBABILIDAD / IMPACTO**")
    for i, label in enumerate(IMPACT_ORDER, start=1):
        header_cols[i].markdown(f"**{label}**")

    colors = {"Bajo": "#dcfce7", "Medio": "#fde68a", "Alto": "#fecaca"}

    for p in PROB_ORDER:
        row_cols = st.columns([1.3, 1, 1, 1, 1, 1])
        row_cols[0].markdown(f"**{p}**")
        for idx, imp in enumerate(IMPACT_ORDER, start=1):
            risk = RISK_MATRIX[(p, imp)]
            secciones = df_matriz.loc[
                (df_matriz["Nivel de probabilidad"] == p) & (df_matriz["Nivel de impacto"] == imp),
                "Sección"
            ].astype(str).tolist()
            content = "<br>".join(html.escape(s) for s in secciones) if secciones else "&nbsp;"
            row_cols[idx].markdown(
                f"""
                <div style="border:1px solid #d1d5db;border-radius:8px;padding:0.55rem;min-height:92px;
                            background:{colors[risk]};font-size:0.9rem;">
                    <div style="font-weight:700;margin-bottom:0.35rem;">{risk}</div>
                    <div>{content}</div>
                </div>
                """,
                unsafe_allow_html=True
            )



# ============================================================
# GENERACIÓN DOCX - ANEXO I PACF
# ============================================================
def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, align="left", size=9, color="000000"):
    cell.text = ""
    p = cell.paragraphs[0]
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("" if text is None else str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def format_df_for_docx(df: pd.DataFrame, percent_cols=None, euro_cols=None, int_cols=None):
    percent_cols = percent_cols or []
    euro_cols = euro_cols or []
    int_cols = int_cols or []
    out = df.copy()
    for c in percent_cols:
        if c in out.columns:
            out[c] = out[c].apply(lambda v: fmt_es_pct(v))
    for c in euro_cols:
        if c in out.columns:
            out[c] = out[c].apply(lambda v: fmt_es_eur(v))
    for c in int_cols:
        if c in out.columns:
            out[c] = out[c].apply(lambda v: fmt_es_num(v, 0))
    return out.fillna("")


def add_dataframe_table(doc, df: pd.DataFrame, title: str, percent_cols=None, euro_cols=None, int_cols=None):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.style = "Heading 2"
    run = p.add_run(title)
    run.bold = True

    show = format_df_for_docx(df, percent_cols=percent_cols, euro_cols=euro_cols, int_cols=int_cols)

    if show.empty:
        doc.add_paragraph("No hay datos disponibles para esta tabla.")
        return

    table = doc.add_table(rows=1, cols=len(show.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    hdr = table.rows[0].cells
    for i, col in enumerate(show.columns):
        set_cell_shading(hdr[i], "1F4E78")
        set_cell_text(hdr[i], col, bold=True, align="center", size=8, color="FFFFFF")

    for _, row in show.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row.tolist()):
            col = show.columns[i]
            align = "left"
            if col in (percent_cols or []) or col in (euro_cols or []) or col in (int_cols or []):
                align = "center"
            set_cell_text(cells[i], val, align=align, size=8)
    doc.add_paragraph()


def add_method_note(doc):
    p = doc.add_paragraph()
    p.style = "Intense Quote" if "Intense Quote" in [s.name for s in doc.styles] else "Normal"
    p.add_run("Nota metodológica — aplicación y revisión\n").bold = True
    p.add_run(
        "La probabilidad se calcula por ejercicio y sección con la fórmula "
        "Ps = [ (Id*s / Its × 65) + (Its / It × 35) ] ÷ 100. "
        "Para el trienio 2023–2025, el impacto utilizado en la matriz final es el impacto financiero "
        "consolidado, sin incorporar la severidad por inexistencia de dato histórico homogéneo. "
        "La probabilidad se consolida con media ponderada 20/30/50 y el impacto con media simple."
    )


def add_risk_matrix_table(doc, df_matriz: pd.DataFrame):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.style = "Heading 2"
    p.add_run("5.1 Matriz de nivel de riesgo").bold = True

    table = doc.add_table(rows=len(PROB_ORDER) + 1, cols=len(IMPACT_ORDER) + 1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    headers = ["PROBABILIDAD / IMPACTO"] + IMPACT_ORDER
    for c, label in enumerate(headers):
        set_cell_shading(table.cell(0, c), "1F4E78")
        set_cell_text(table.cell(0, c), label, bold=True, align="center", size=8, color="FFFFFF")

    risk_fill = {"Bajo": "DCEED4", "Medio": "F8EDB2", "Alto": "F4CCCC"}

    for r, prob in enumerate(PROB_ORDER, start=1):
        set_cell_shading(table.cell(r, 0), "EAF1F7")
        set_cell_text(table.cell(r, 0), prob, bold=True, align="center", size=8)
        for c, impact in enumerate(IMPACT_ORDER, start=1):
            risk = RISK_MATRIX[(prob, impact)]
            sections = df_matriz.loc[
                (df_matriz["Nivel de probabilidad"] == prob) &
                (df_matriz["Nivel de impacto"] == impact),
                "Sección"
            ].astype(str).tolist()
            content = risk
            if sections:
                content += "\n" + ", ".join(sections)
            set_cell_shading(table.cell(r, c), risk_fill.get(risk, "FFFFFF"))
            set_cell_text(table.cell(r, c), content, bold=True, align="center", size=8)


def build_annex_docx(
    df_summary,
    prob_tables,
    impact_tables,
    prob_final,
    impact_final,
    matriz_final,
):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(2)
    sec.right_margin = Cm(2)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PLAN ANUAL DE CONTROL FINANCIERO 2026")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor.from_string("1F1F1F")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ANEXO I: EVALUACIÓN ESTADÍSTICA DE LOS RIESGOS")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor.from_string("1F4E78")

    doc.add_paragraph(
        "Documento generado automáticamente a partir de los cálculos realizados en la aplicación PACF 2026."
    )

    add_dataframe_table(
        doc,
        df_summary,
        "1. Resumen anual de carga y depuración",
        euro_cols=["Importe total"],
        int_cols=["Registros válidos", "Registros excluidos", "Expedientes únicos", "Secciones"],
    )

    for year in sorted(prob_tables.keys()):
        prob = prob_tables[year]
        impact = impact_tables[year]

        add_dataframe_table(
            doc,
            prob[["Sección", "Descripción", "Id*s", "Its", "P1 (Id*s / Its)", "P2 (Its / It)", "Ps (%)", "Nivel de probabilidad"]],
            f"2.{year} Probabilidad del riesgo de incumplimiento — Ejercicio {year}",
            percent_cols=["P1 (Id*s / Its)", "P2 (Its / It)", "Ps (%)"],
            int_cols=["Id*s", "Its"],
        )

        add_dataframe_table(
            doc,
            impact[["Sección", "Descripción", "Ms", "M", "Is (%)", "Nivel de impacto"]],
            f"3.{year} Impacto económico — Ejercicio {year}",
            euro_cols=["Ms", "M"],
            percent_cols=["Is (%)"],
        )

    col_media_prob_docx = "Media ponderada probabilidad" if "Media ponderada probabilidad" in prob_final.columns else next(
        (c for c in prob_final.columns if c.startswith("Media ponderada")),
        None
    )
    cols_prob_docx = ["Sección", "Descripción"] + [c for c in prob_final.columns if c.startswith("Ps ")]
    if col_media_prob_docx:
        cols_prob_docx.append(col_media_prob_docx)
    cols_prob_docx.append("Nivel de probabilidad")
    cols_prob_docx = [c for c in cols_prob_docx if c in prob_final.columns]

    pct_prob_docx = [c for c in prob_final.columns if c.startswith("Ps ")]
    if col_media_prob_docx and col_media_prob_docx in prob_final.columns:
        pct_prob_docx.append(col_media_prob_docx)

    add_dataframe_table(
        doc,
        prob_final[cols_prob_docx],
        "4.1 Media ponderada de probabilidad",
        percent_cols=pct_prob_docx,
    )

    add_dataframe_table(
        doc,
        impact_final[["Sección", "Descripción"] + [c for c in impact_final.columns if c.startswith("Is ")] + ["Media trienal impacto", "Nivel de impacto", "Severidad", "Modo impacto"]],
        "4.2 Media trienal de impacto",
        percent_cols=[c for c in impact_final.columns if c.startswith("Is ")] + ["Media trienal impacto"],
    )

    col_media_matriz_docx = "Media ponderada probabilidad" if "Media ponderada probabilidad" in matriz_final.columns else next(
        (c for c in matriz_final.columns if c.startswith("Media ponderada")),
        None
    )
    cols_matriz_docx = ["Sección", "Descripción"]
    if col_media_matriz_docx:
        cols_matriz_docx.append(col_media_matriz_docx)
    cols_matriz_docx += ["Nivel de probabilidad", "Media trienal impacto", "Nivel de impacto", "Nivel de riesgo"]
    cols_matriz_docx = [c for c in cols_matriz_docx if c in matriz_final.columns]

    pct_matriz_docx = [c for c in [col_media_matriz_docx, "Media trienal impacto"] if c and c in matriz_final.columns]

    add_dataframe_table(
        doc,
        matriz_final[cols_matriz_docx],
        "5. Mapa del riesgo resultante — PACF 2026",
        percent_cols=pct_matriz_docx,
    )

    add_risk_matrix_table(doc, matriz_final)
    add_method_note(doc)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()

# ============================================================
# CARGA DE ARCHIVOS
# ============================================================
with st.sidebar:
    st.header("Carga de datos")
    uploaded_map = st.file_uploader("Excel de mapeo de secciones", type=["xls", "xlsx"], key="map")

    st.subheader("Parámetros de ponderación")
    st.caption("Pesos para la media de probabilidad de los ejercicios cargados.")
    peso_1 = st.number_input("Peso ejercicio más antiguo (%)", min_value=0.0, max_value=100.0, value=20.0, step=1.0, key="peso1")
    peso_2 = st.number_input("Peso ejercicio intermedio (%)", min_value=0.0, max_value=100.0, value=30.0, step=1.0, key="peso2")
    peso_3 = st.number_input("Peso ejercicio más reciente (%)", min_value=0.0, max_value=100.0, value=50.0, step=1.0, key="peso3")

    st.subheader("Modo de cálculo de Ms")
    modo_ms = st.radio(
        "Selecciona el criterio de impacto financiero",
        options=[
            "Todos los informes válidos",
            "Solo peticiones con 1 o más desfavorables",
        ],
        index=1,
        key="modo_ms"
    )

    years_input = []
    uploaded_inputs = []
    for idx, default_year in enumerate([2023, 2024, 2025], start=1):
        st.subheader(f"Ejercicio {idx}")
        year_val = st.number_input(f"Año {idx}", min_value=2000, max_value=2100, value=default_year, step=1, key=f"year_{idx}")
        file_val = st.file_uploader(f"Excel ejercicio {year_val}", type=["xls", "xlsx"], key=f"file_{idx}")
        years_input.append(int(year_val))
        uploaded_inputs.append(file_val)

uploaded_years = [year for year, uploaded in zip(years_input, uploaded_inputs) if uploaded is not None]
if len(uploaded_years) != len(set(uploaded_years)):
    st.error("Hay ejercicios duplicados con archivo cargado. Revisa los años para que cada Excel anual tenga un ejercicio distinto.")
    st.stop()

loaded_raw = {}
for year, uploaded in zip(years_input, uploaded_inputs):
    if uploaded is not None:
        loaded_raw[year] = sanitize_columns(leer_excel(uploaded))

if not loaded_raw:
    st.info("Sube al menos un Excel anual y, si lo tienes, el mapeo de secciones.")
    st.stop()

df_map = None
if uploaded_map is not None:
    try:
        df_map = preparar_mapeo_secciones(leer_excel(uploaded_map))
        st.success("Mapeo de secciones cargado correctamente.")
    except Exception as e:
        st.error(f"Error en el mapeo de secciones: {e}")
        st.stop()

# ============================================================
# MAPEADO DE COLUMNAS
# ============================================================
st.markdown('<div class="section-card"><b>1. Mapeo de columnas</b><br>Selecciona los campos reales de cada Excel. La aplicación normalizará la estructura para producir las tablas del anexo.</div>', unsafe_allow_html=True)

normalized_by_year = {}
excluded_by_year = {}
frames_export = {}

for year in sorted(loaded_raw.keys()):
    df_raw = loaded_raw[year]
    st.markdown(f"#### Ejercicio {year}")
    with st.expander(f"Columnas detectadas en {year}", expanded=False):
        st.write(df_raw.columns.tolist())
        st.dataframe(df_raw.head(8), use_container_width=True, hide_index=True)

    cols = [""] + df_raw.columns.tolist()
    c1, c2, c3 = st.columns(3)
    with c1:
        expediente_col = selectbox_column(
            f"{year} · Nº Expediente FLP",
            cols,
            aliases={"Nº Expediente FLP", "Numero Expediente FLP", "Número Expediente FLP", "Expediente", "N Expediente"},
            contains_all=["expediente"],
            key=f"exp_{year}",
        )
        importe_col = selectbox_column(
            f"{year} · Importe",
            cols,
            aliases={"Importe", "Importe total", "Importe expediente", "Cuantía", "Cuantia"},
            contains_all=["importe"],
            key=f"imp_{year}",
        )
    with c2:
        desf_col = selectbox_column(
            f"{year} · Informes desfavorables",
            cols,
            aliases={"Número de Informes Desfavorables", "Numero de Informes Desfavorables", "Informes desfavorables", "Desfavorables"},
            contains_all=["desfavorables"],
            key=f"desf_{year}",
        )
        fav_col = selectbox_column(
            f"{year} · Informes favorables",
            cols,
            aliases={"Número de Informes Favorables", "Numero de Informes Favorables", "Informes favorables", "Favorables"},
            contains_all=["favorables"],
            key=f"fav_{year}",
        )
    with c3:
        estado_col = selectbox_column(
            f"{year} · Estado",
            cols,
            aliases={"Estado", "Situación", "Situacion", "Estado expediente"},
            contains_all=["estado"],
            key=f"estado_{year}",
        )
        fase_col = selectbox_column(
            f"{year} · Fase del gasto",
            cols,
            aliases={"Fase del Gasto", "Fase gasto", "Fase", "Tipo de fase"},
            contains_all=["fase"],
            key=f"fase_{year}",
        )

    col_map = {
        "expediente": expediente_col or None,
        "importe": importe_col or None,
        "desfavorables": desf_col or None,
        "favorables": fav_col or None,
        "estado": estado_col or None,
        "fase": fase_col or None,
    }

    try:
        df_norm = normalize_year_df(df_raw, year, col_map)
        df_ok, df_exc = depurar_df(df_norm)
        df_ok = aplicar_mapeo(df_ok, df_map)
        df_exc = aplicar_mapeo(df_exc, df_map) if not df_exc.empty else df_exc

        normalized_by_year[year] = df_ok
        excluded_by_year[year] = df_exc

        st.success(f"Ejercicio {year} preparado correctamente. Registros válidos: {len(df_ok)}")
    except Exception as e:
        st.error(f"No se ha podido preparar el ejercicio {year}: {e}")

if not normalized_by_year:
    st.warning("No hay ejercicios válidos para continuar.")
    st.stop()

# ============================================================
# RESUMEN
# ============================================================
st.markdown('<div class="section-card"><b>2. Resumen de carga y depuración</b></div>', unsafe_allow_html=True)

summary_rows = []
for year in sorted(normalized_by_year.keys()):
    df_ok = normalized_by_year[year]
    df_exc = excluded_by_year.get(year, pd.DataFrame())
    summary_rows.append({
        "Año": year,
        "Registros válidos": len(df_ok),
        "Registros excluidos": len(df_exc),
        "Expedientes únicos": df_ok["Nº Expediente FLP"].nunique(),
        "Secciones": df_ok["Sección"].nunique(),
        "Importe total": df_ok["Importe"].sum()
    })

df_summary = pd.DataFrame(summary_rows)
frames_export["00_Resumen"] = df_summary

c1, c2, c3, c4 = st.columns(4)
c1.markdown(f'<div class="kpi"><div class="kpi-title">Ejercicios cargados</div><div class="kpi-value">{len(normalized_by_year)}</div></div>', unsafe_allow_html=True)
c2.markdown(f'<div class="kpi"><div class="kpi-title">Registros válidos</div><div class="kpi-value">{fmt_es_num(df_summary["Registros válidos"].sum(),0)}</div></div>', unsafe_allow_html=True)
c3.markdown(f'<div class="kpi"><div class="kpi-title">Registros excluidos</div><div class="kpi-value">{fmt_es_num(df_summary["Registros excluidos"].sum(),0)}</div></div>', unsafe_allow_html=True)
c4.markdown(f'<div class="kpi"><div class="kpi-title">Importe agregado</div><div class="kpi-value">{fmt_es_eur(df_summary["Importe total"].sum())}</div></div>', unsafe_allow_html=True)

display_visual_table(df_summary, "Resumen anual de carga", euro_cols=["Importe total"], int_cols=["Registros válidos", "Registros excluidos", "Expedientes únicos", "Secciones"])

# ============================================================
# TABLAS ANUALES
# ============================================================
st.markdown('<div class="section-card"><b>3. Evaluación anual por ejercicio</b></div>', unsafe_allow_html=True)

prob_tables = {}
impact_tables = {}

for year in sorted(normalized_by_year.keys()):
    df_year = normalized_by_year[year]
    st.markdown(f"## Ejercicio {year}")

    prob = calcular_probabilidad_anual(df_year, df_map)
    impact = calcular_impacto_anual(df_year, df_map, modo_ms=modo_ms)
    prob_tables[year] = prob
    impact_tables[year] = impact

    frames_export[f"Probabilidad_{year}"] = prob
    frames_export[f"Impacto_{year}"] = impact
    frames_export[f"Detalle_{year}"] = df_year
    frames_export[f"Excluidos_{year}"] = excluded_by_year.get(year, pd.DataFrame())

    tab1, tab2, tab3 = st.tabs(["Probabilidad", "Impacto", "Detalle"])

    with tab1:
        display_visual_table(
            prob[["Sección", "Descripción", "Id*s", "Its", "P1 (Id*s / Its)", "P2 (Its / It)", "Ps (%)", "Nivel de probabilidad"]],
            f"Probabilidad del riesgo de incumplimiento — {year}",
            percent_cols=["P1 (Id*s / Its)", "P2 (Its / It)", "Ps (%)"],
            int_cols=["Id*s", "Its"],
            badge_cols=["Nivel de probabilidad"]
        )

    with tab2:
        display_visual_table(
            impact[["Sección", "Descripción", "Ms", "M", "Is (%)", "Nivel de impacto"]],
            f"Impacto económico — {year}",
            euro_cols=["Ms", "M"],
            percent_cols=["Is (%)"],
            badge_cols=["Nivel de impacto"]
        )

    with tab3:
        detail_cols = [c for c in ["Nº Expediente FLP", "Sección", "Descripción", "Fase del Gasto", "Importe", "Número de Informes Favorables", "Número de Informes Desfavorables", "Estado"] if c in df_year.columns]
        display_visual_table(
            df_year[detail_cols],
            f"Detalle de expedientes — {year}",
            euro_cols=["Importe"],
            int_cols=["Número de Informes Favorables", "Número de Informes Desfavorables"]
        )

# ============================================================
# CONSOLIDACIÓN
# ============================================================
st.markdown('<div class="section-card"><b>4. Consolidación trienal</b></div>', unsafe_allow_html=True)

pesos_probabilidad = [peso_1, peso_2, peso_3]
prob_final = consolidar_probabilidad(prob_tables, df_map, pesos_probabilidad=pesos_probabilidad)
impact_final = consolidar_impacto(impact_tables, df_map)
matriz_final = construir_matriz_final(prob_final, impact_final)

frames_export["Probabilidad_final"] = prob_final
frames_export["Impacto_final"] = impact_final
frames_export["Matriz_final"] = matriz_final

tabf1, tabf2, tabf3 = st.tabs(["Probabilidad final", "Impacto final", "Matriz de riesgo"])

with tabf1:
    col_media_prob = next((c for c in prob_final.columns if c.startswith("Media ponderada (")), "Media ponderada probabilidad")
    pcols = [c for c in prob_final.columns if c.startswith("Ps ")] + [col_media_prob]
    display_visual_table(
        prob_final[["Sección", "Descripción"] + [c for c in prob_final.columns if c.startswith("Ps ")] + [col_media_prob, "Nivel de probabilidad", "Pesos aplicados"]],
        "Media ponderada de probabilidad",
        percent_cols=pcols,
        badge_cols=["Nivel de probabilidad"]
    )

with tabf2:
    icols = [c for c in impact_final.columns if c.startswith("Is ")] + ["Media trienal impacto"]
    display_visual_table(
        impact_final[["Sección", "Descripción"] + [c for c in impact_final.columns if c.startswith("Is ")] + ["Media trienal impacto", "Nivel de impacto", "Severidad", "Modo impacto"]],
        "Media trienal de impacto",
        percent_cols=icols,
        badge_cols=["Nivel de impacto"]
    )

with tabf3:
    cols_matriz = ["Sección", "Descripción", "Media ponderada probabilidad", "Nivel de probabilidad", "Media trienal impacto", "Nivel de impacto", "Nivel de riesgo"]
    cols_matriz = [c for c in cols_matriz if c in matriz_final.columns]
    pct_matriz = [c for c in ["Media ponderada probabilidad", "Media trienal impacto"] if c in matriz_final.columns]
    badge_matriz = [c for c in ["Nivel de probabilidad", "Nivel de impacto", "Nivel de riesgo"] if c in matriz_final.columns]
    display_visual_table(
        matriz_final[cols_matriz],
        "Mapa del riesgo resultante",
        percent_cols=pct_matriz,
        badge_cols=badge_matriz
    )
    display_matrix_grid(matriz_final)

# ============================================================
# NOTA METODOLÓGICA
# ============================================================
st.markdown('<div class="section-card"><b>5. Nota metodológica</b></div>', unsafe_allow_html=True)
st.markdown("""
<div class="method-note">
<b>Probabilidad.</b> Se calcula por ejercicio y sección con la fórmula:
<i>Ps = [ (Id*s / Its × 65) + (Its / It × 35) ] ÷ 100</i>.<br><br>
<b>Impacto.</b> Para el trienio 2023–2025 se utiliza el <b>impacto financiero</b>.<br>
El cálculo de <b>Ms</b> es <b>configurable</b> desde la aplicación: puede hacerse con <b>todos los informes válidos</b> o solo con las <b>peticiones que tienen 1 o más informes desfavorables</b>. En ambos casos, <b>M</b> se mantiene como el importe global de todos los informes válidos del ejercicio.<br>
La <b>severidad</b> queda preparada en la aplicación, pero no se incorpora al cálculo final por inexistencia de dato histórico homogéneo en las hojas de entrada.<br><br>
<b>Consolidación.</b> La probabilidad se consolida con una <b>media ponderada parametrizable</b> desde la aplicación.
El impacto se consolida con media simple de los ejercicios cargados.
</div>
""", unsafe_allow_html=True)

# ============================================================
# DESCARGA
# ============================================================
st.markdown('<div class="section-card"><b>6. Exportación</b></div>', unsafe_allow_html=True)
excel_bytes = build_excel_export(frames_export)
st.download_button(
    "⬇️ Descargar resultados en Excel",
    data=excel_bytes,
    file_name="PACF_2026_anexo_visual_resultados.xlsx",
    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
)

docx_bytes = build_annex_docx(
    df_summary=df_summary,
    prob_tables=prob_tables,
    impact_tables=impact_tables,
    prob_final=prob_final,
    impact_final=impact_final,
    matriz_final=matriz_final,
)
st.download_button(
    "⬇️ Descargar Anexo I en Word",
    data=docx_bytes,
    file_name="Anexo_I_PACF_2026.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
